import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_report():
    doc = Document()
    
    # Page Setup - Normal Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # --- PORTADA / ENCABEZADO ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    
    run_badge = title_p.add_run("CONTROL DE BOLSAS — MASTER TRACK ERP\n")
    run_badge.font.size = Pt(12)
    run_badge.font.bold = True
    run_badge.font.color.rgb = RGBColor(37, 99, 235) # Royal Blue
    
    run_title = title_p.add_run("REPORTE TÉCNICO DE ARQUITECTURA,\nFUNCIONES Y MEJORAS DEL SISTEMA")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("Versión del Sistema: v8.9.53 Enterprise | Fecha de Auditoría: Agosto 2026\nPlataforma: React 18 + Vite + TypeScript + Firebase Cloud Platform")
    run_sub.font.size = Pt(10.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- RESUMEN EJECUTIVO BOX ---
    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = summary_table.cell(0, 0)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p_box = cell.paragraphs[0]
    p_box.paragraph_format.space_after = Pt(4)
    r_box_t = p_box.add_run("📌 Resumen Ejecutivo de Operación:\n")
    r_box_t.font.bold = True
    r_box_t.font.size = Pt(11)
    r_box_t.font.color.rgb = RGBColor(15, 23, 42)
    
    r_box_c = p_box.add_run(
        "El presente documento detalla la arquitectura de software, especificación de módulos, fórmulas matemáticas de liquidación, "
        "reglas de negocio y mecanismos de seguridad del ERP 'Control de Bolsas — Master Track'. El sistema gestiona integralmente "
        "el flujo de compra de polietileno a maquila (Andrés), entrega en báscula, facturación electrónica CFDI 4.0 a Grupo Textil Providencia "
        "(Textil Hogar y Planta P4), cobranza de contrarecibos, tesorería de caja chica y reparto de utilidades 50/50."
    )
    r_box_c.font.size = Pt(10)
    r_box_c.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # --- SECCIÓN 1: ARQUITECTURA GENERAL ---
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Arquitectura Técnica de la Plataforma")
    r_h1.font.color.rgb = RGBColor(15, 23, 42)
    
    p_arch = doc.add_paragraph(
        "El sistema opera bajo un modelo modular desacoplado en la nube con capacidad de funcionamiento sin conexión (Progressive Web App):"
    )
    p_arch.paragraph_format.space_after = Pt(8)
    
    arch_table = doc.add_table(rows=5, cols=2)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch_data = [
        ("Capa de Software", "Tecnología & Componentes Clave"),
        ("Frontend (Cliente)", "React 18, Vite 5, TypeScript 5, Framer Motion, Vanilla CSS Custom Design System, Service Workers (PWA Offline)."),
        ("Motor Lógico & Parsers", "decimal.js-light (aritmética de punto fijo), xmlParser (CFDI 4.0 SAT), pdfmake (generador PDF), cryptoAudit (SHA-256)."),
        ("Backend & Serverless", "Firebase Cloud Functions v2 (Node.js 22), Gemini 2.5 Flash Multimodal (extracción IA de remisiones de báscula)."),
        ("Persistencia & Seguridad", "Cloud Firestore (NoSQL transaccional), Firebase Storage, Firebase Authentication (RBAC: Admin, Manager, Viewer).")
    ]
    
    for row_idx, row in enumerate(arch_table.rows):
        for col_idx, text in enumerate(arch_data[row_idx]):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            if row_idx == 0:
                set_cell_background(cell, "1E293B")
                p = cell.paragraphs[0]
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            else:
                set_cell_background(cell, "F8FAFC" if row_idx % 2 == 1 else "FFFFFF")
                p = cell.paragraphs[0]
                p.runs[0].font.size = Pt(10)
                if col_idx == 0:
                    p.runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # --- SECCIÓN 2: DESGLOSE TÉCNICO DE CADA FUNCIÓN DEL ERP ---
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. Desglose Técnico Función por Función")
    r_h2.font.color.rgb = RGBColor(15, 23, 42)

    # Función A
    h_a = doc.add_heading(level=2)
    h_a.add_run("A. Módulo de Operación y Expedientes Maestros (/ordenes y OrderModal)")
    doc.add_paragraph(
        "• Pipeline de 6 Estaciones (OrderStepper): Controla el ciclo de vida de cada Orden de Compra desde su emisión hasta su liquidación:\n"
        "   1. Captura de OC  ➔  2. Extrusión Andrés  ➔  3. Entregas Báscula  ➔  4. Facturación CFDI  ➔  5. Cobranza Providencia  ➔  6. Liquidación 50/50.\n"
        "• Cierre Forzado por Menos Kilos (isClosedShort): Cuando el proveedor entrega un volumen inferior al estipulado en la OC pero concluye el pedido, "
        "el usuario activa este botón en 1 clic. El sistema elimina las alertas de kilos pendientes y permite facturar y liquidar el 100% de lo entregado.\n"
        "• Separación Departamental Estricta: Textil Hogar (TH - Nava / José Antonio Torre Lamuño) con prefijo TH- (ej. TH-1024) y Grupo Textil Providencia "
        "(GT - Evelia / Planta P4) con prefijo GT- (ej. GT-890). Cero mezcla de entregas ni facturas entre departamentos."
    )

    # Función B
    h_b = doc.add_heading(level=2)
    h_b.add_run("B. Motor Matemático y Fórmulas Financieras (src/lib/finance.ts)")
    doc.add_paragraph(
        "Toda operación contable opera bajo aritmética de precisión absoluta sin flotantes IEEE 754 mediante las siguientes ecuaciones oficiales:\n\n"
        "1. Facturación a Providencia:\n"
        "   • Subtotal Venta = Kilos Facturados × $43.00 MXN\n"
        "   • Total con IVA (16%) = Subtotal Venta × 1.16 = Kilos Facturados × $49.88 MXN\n\n"
        "2. Costo de Compra a Andrés (Maquila):\n"
        "   • Costo Andrés = Kilos Facturados × $38.00 MXN\n\n"
        "3. Comisión Contable:\n"
        "   • Comisión Contador = Subtotal Venta × 8.00%\n\n"
        "4. Margen Bruto de Operación:\n"
        "   • Margen Bruto = $43.00 - $38.00 = $5.00 MXN por kilogramo\n\n"
        "5. Utilidad Líquida y Reparto de Socios (P&L 50/50):\n"
        "   • Utilidad Real = Subtotal Facturado - Costo Andrés ($38/kg) - Comisión Contador (8%) - Gastos Operativos\n"
        "   • Reparto Paco (50%) = Utilidad Real / 2\n"
        "   • Reparto Socio (50%) = Utilidad Real / 2"
    )

    # Función C
    h_c = doc.add_heading(level=2)
    h_c.add_run("C. Módulo de Proveedor / Maquila y Libro Mayor de Andrés (/compras)")
    doc.add_paragraph(
        "• Convención de Signos del Libro Mayor:\n"
        "   - Saldo Positivo (+): Anticipo de la empresa a favor en manos de Andrés (saldo a favor pendiente de amortizar con material).\n"
        "   - Saldo Negativo (-): Deuda o pasivo de la empresa hacia Andrés.\n"
        "• Fórmula de Saldo Vivo:\n"
        "   Saldo Andrés = Saldo Histórico + Pagos/Anticipos - (Kilos Entregados Canónicos × $38.00)\n"
        "• Blindaje Canónico Inviolable: Total de entregas fijado en exactamente 10,366.21 kg ($393,915.98 MXN), garantizando que registros "
        "antiguos de contrarecibos no contaminen el kardex de materia prima."
    )

    # Función D
    h_d = doc.add_heading(level=2)
    h_d.add_run("D. Módulo de Cobranza y Cartera Providencia (/cobranza)")
    doc.add_paragraph(
        "• Cardinalidad de Contrarecibos (1 a N): Un Contrarecibo oficial de Providencia ampara una o varias facturas; una factura solo puede "
        "estar vinculada a un único Contrarecibo.\n"
        "• Semáforo de Vencimiento y Mora: El sistema calcula la fecha límite a 30 días naturales posteriores a la emisión del Contrarecibo. "
        "Una factura nunca se marca como vencida si aún no cuenta con Contrarecibo emitido por Providencia.\n"
        "• Estado de Cuenta Espejo en PDF: Generador de reportes con membrete institucional idéntico al formato interno de Providencia para conciliación inmediata."
    )

    # Función E
    h_e = doc.add_heading(level=2)
    h_e.add_run("E. Módulo de Caja Chica y Tesorería en Efectivo (/caja-chica)")
    doc.add_paragraph(
        "• Control de Ingresos y Egresos: Registro transaccional en tiempo real con cálculo inmutable de saldo disponible en mano.\n"
        "• Clasificación Contable: Gastos operativos, viáticos, pagos de maniobra de chofer, amortizaciones a Andrés y traspasos de utilidades.\n"
        "• Exportación de Cortes Semanales y Mensuales: Auditoría de balance con generación de PDF y Excel."
    )

    # Función F
    h_f = doc.add_heading(level=2)
    h_f.add_run("F. Parsers Inteligentes, Gemini AI y Motor de Seguridad (src/lib/ & functions/)")
    doc.add_paragraph(
        "• Parser XML CFDI 4.0: Extracción instantánea de UUID fiscal, RFC emisor/receptor, conceptos, importes y validación de sellos SAT.\n"
        "• Extractor Multimodal Gemini 2.5 Flash: Lectura de fotos de tickets de báscula y notas de remisión para carga automática de kilos.\n"
        "• Guardas Anti-Duplicados (duplicateGuards.ts): Bloqueo criptográfico ante intentos de capturar folios repetidos de facturas o contrarecibos.\n"
        "• Suite de Pruebas Unitarias: 124 tests automáticos (npm test) que validan el 100% de la lógica contable antes de cada compilación."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # --- SECCIÓN 3: PROPUESTAS DE MEJORA ---
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Propuestas de Mejora e Innovación para el ERP")
    r_h3.font.color.rgb = RGBColor(15, 23, 42)
    
    doc.add_paragraph(
        "Tomando como base la sólida arquitectura actual, se presentan 6 mejoras de alto impacto para maximizar la automatización y eficiencia del negocio:"
    )

    mejoras = [
        ("1. Conciliación Bancaria Automática (OCR / Excel Bancario)",
         "Permitir subir el estado de cuenta bancario (BBVA, Banamex, etc.) en PDF o Excel. Un algoritmo inteligente cruza los depósitos con los contrarecibos pendientes y los marca como 'Cobrados' en un solo clic, ahorrando 40 minutos de conciliación manual semanal."),
        
        ("2. Notificaciones y Alertas Proactivas vía WhatsApp / Telegram Bot",
         "Implementar un bot automatizado en Cloud Functions que envíe diariamente a las 8:00 AM un resumen con: contrarecibos que vencen hoy, facturas en revisión y entregas en patio pendientes de facturar."),
        
        ("3. Pronóstico de Flujo de Efectivo a 30/60 Días (Cash Flow Forecasting)",
         "Una gráfica interactiva en el Dashboard que proyecte las entradas de cobranza esperadas según las fechas de vencimiento de los contrarecibos frente a los compromisos de pago a Andrés y gastos fijos."),
        
        ("4. Remisiones de Báscula con Código QR Inteligente",
         "Incorporar un código QR en las remisiones impresas. Al escanearlo con el celular del chofer o en patio, el sistema carga automáticamente la entrega sin necesidad de escribir manualmente el folio de la OC."),
        
        ("5. Simulador de Negociación y Escenarios de Margen (What-If Analysis)",
         "Herramienta interactiva para calcular en segundos el impacto en el reparto de utilidades si cambian los precios (ej. si Andrés ajusta el costo por kilo o Providencia renegocia el precio de venta)."),
         
        ("6. Respaldo Automático a la Nube (Google Drive / Dropbox)",
         "Además del script USB local, programar un respaldo encriptado diario automático hacia una carpeta segura de Google Drive corporativo.")
    ]

    for titulo, desc in mejoras:
        p_m = doc.add_paragraph()
        p_m.paragraph_format.space_after = Pt(6)
        r_t = p_m.add_run(f"⭐ {titulo}\n")
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = RGBColor(37, 99, 235)
        
        r_d = p_m.add_run(desc)
        r_d.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # --- FIRMA / PIE ---
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_foot = p_footer.add_run("Generado automáticamente por Antigravity AI Engine\nControl de Bolsas — Master Track ERP v8.9.53")
    r_foot.font.size = Pt(9.5)
    r_foot.font.italic = True
    r_foot.font.color.rgb = RGBColor(148, 163, 184)

    # Guardar en c:\pacoputo
    output_docx = r"c:\pacoputo\REPORTE_TECNICO_ERP_CONTROL_DE_BOLSAS.docx"
    output_doc = r"c:\pacoputo\REPORTE_TECNICO_ERP_CONTROL_DE_BOLSAS.doc"
    doc.save(output_docx)
    doc.save(output_doc)
    print(f"Reporte generado exitosamente en:\n- {output_docx}\n- {output_doc}")

if __name__ == '__main__':
    create_report()
